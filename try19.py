import pandas as pd
import numpy as np
import csv
import matplotlib.pyplot as plt
from datetime import date
import mysql.connector as s
import docx
uf=pd.read_csv('History.csv')
usn=input("\t\t\t\t Enter User Name:-")
uf1=uf[(uf['User Name']==usn)]
if(uf1.empty):
    print("\t\t\t           OOPS! User not found")
    print("\t\t< Weather You entered Wrong Name or User Not Present >")
else:
    data=(uf[uf['User Name']==usn].index.values)
    print(uf.loc[data])
    us=int(input("\t\t\t\t Enter The Transection Id:-"))
    uf2=uf[(uf['User Name']==usn)&(uf['Transection ID']==us)]
    if(uf2.empty):
        print("\t\t\t           OOPS! Transection not found")
        print("\t\t< Weather You entered ID >")
    else:
        dat=(uf[uf['Transection ID']==us].index.values)
        doc=docx.Document()
                        doc.add_paragraph("~:~"*23)
                        a=doc.add_paragraph("S.A.S Online Transection")
                        a.alignment=1
                        doc.add_paragraph("~:~"*23)
                        ab="Transection ID:~"+str(us)
                        a=doc.add_paragraph(ab)
                        a.alignment=1
                        b=np.array(uf)
                        usn=b[dat,1]
                        usr=usn
                        ab="Transection Done Using User Name:~"+usr
                        a=doc.add_paragraph(ab)
                        a.alignment=1
                        usn=b[data,3]
                        amu=usn
                        ab="Transection Type:~"+ amu
                        a=doc.add_paragraph(ab)
                        a.alignment=1
                        usn=b[data,4]
                        no=usn
                        ab="Number as per Type:~"+str(*no)
                        a=doc.add_paragraph(ab)
                        a.alignment=1
                        usn=b[data,2]
                        tm=usn
                        ab="Amount:~"+str(*tm)
                        a=doc.add_paragraph(ab)
                        a.alignment=1
                        usn=b[data,5]
                        do=usn
                        ab="Transection Date:~"+str(*do)
                        a=doc.add_paragraph(ab)
                        a.alignment=1
                        a=doc.add_paragraph("|| Successfully Paid ||")
                        a.alignment=1
                        a=doc.add_paragraph("||* Thanks ! For Using *||")
                        a.alignment=1
                        doc.add_paragraph("~:~"*23)
                        doc.save("Recepit.docx")
                        d=docx.Document("Recepit.docx")
                        pa=d.paragraphs
                        for para in pa:
                                print(para.text)
                        p=1
                        q=0
                        w=0
                        subprocess.Popen("Recepit.docx")
